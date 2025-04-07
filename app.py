import streamlit as st
import os
import tempfile
import subprocess
import json
import time
import datetime
import io
from pathlib import Path
from typing import List, Dict, Any, Optional
import logging
from openai import OpenAI
import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Setup logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# Page configuration
st.set_page_config(
    page_title="Система помощника следователя",
    page_icon="🔍",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS
def load_css():
    st.markdown("""
        <style>
        /* Main styles */
        .stApp {
            background-color: #F3F4F6;
        }
        
        /* Cards */
        .card {
            padding: 1.5rem;
            border-radius: 0.5rem;
            background-color: white;
            box-shadow: 0 1px 3px rgba(0,0,0,0.1);
            margin-bottom: 1rem;
        }
        
        .card-header {
            font-weight: bold;
            font-size: 1.2rem;
            margin-bottom: 1rem;
            padding-bottom: 0.5rem;
            border-bottom: 1px solid #e5e7eb;
        }
        
        /* Buttons */
        .primary-btn {
            background-color: #1E40AF;
            color: white;
            padding: 0.5rem 1rem;
            border-radius: 0.375rem;
            border: none;
            cursor: pointer;
            font-weight: 500;
            width: 100%;
        }
        
        .primary-btn:hover {
            background-color: #1E3A8A;
        }
        
        /* Form elements */
        .form-input {
            padding: 0.5rem;
            border: 1px solid #D1D5DB;
            border-radius: 0.375rem;
            width: 100%;
            margin-bottom: 0.5rem;
        }
        
        /* Tables */
        .styled-table {
            width: 100%;
            border-collapse: collapse;
        }
        
        .styled-table th, .styled-table td {
            padding: 0.75rem;
            text-align: left;
            border-bottom: 1px solid #e5e7eb;
        }
        
        .styled-table th {
            background-color: #f9fafb;
            font-weight: 500;
        }
        
        .footer {
            text-align: center;
            padding: 1rem;
            color: #6B7280;
            font-size: 0.875rem;
            margin-top: 2rem;
        }
        </style>
    """, unsafe_allow_html=True)

# Initialize OpenAI client
def init_openai():
    api_key = st.secrets.get("openai_api_key")
    if not api_key:
        st.error("API ключ OpenAI не настроен. Пожалуйста, добавьте его в настройки секретов.")
        return None
    
    try:
        return OpenAI(api_key=api_key)
    except Exception as e:
        st.error(f"Ошибка инициализации OpenAI API: {str(e)}")
        return None

# Create necessary directories
def create_directories():
    try:
        Path("storage/transcriptions").mkdir(parents=True, exist_ok=True)
        Path("storage/planning").mkdir(parents=True, exist_ok=True)
        Path("storage/indictments").mkdir(parents=True, exist_ok=True)
        Path("storage/methodologies").mkdir(parents=True, exist_ok=True)
        Path("storage/evidence").mkdir(parents=True, exist_ok=True)
        Path("storage/temp").mkdir(parents=True, exist_ok=True)
    except Exception as e:
        st.error(f"Ошибка при создании директорий: {str(e)}")

# Check for FFmpeg availability
def check_ffmpeg():
    try:
        subprocess.run(["ffmpeg", "-version"], check=True, capture_output=True)
        return True
    except (subprocess.CalledProcessError, FileNotFoundError):
        return False

# Generate unique case number
def generate_case_number(prefix="М"):
    now = datetime.datetime.now()
    return f"{prefix}-{now.strftime('%Y%m%d-%H%M%S')}"

# Save session history
def save_history(module_type, data):
    """Save history data to the appropriate directory"""
    create_directories()
    
    filename = f"{module_type}_{data.get('id', datetime.datetime.now().strftime('%Y%m%d%H%M%S'))}.json"
    filepath = f"storage/{module_type}/{filename}"
    
    with open(filepath, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)
    
    return filepath

# Load session history
def load_history(module_type):
    """Load all history data from the specified module directory"""
    history_path = f"storage/{module_type}"
    if not os.path.exists(history_path):
        return []
    
    history_files = [f for f in os.listdir(history_path) if f.endswith('.json')]
    history = []
    
    for file in history_files:
        try:
            with open(os.path.join(history_path, file), "r", encoding="utf-8") as f:
                history.append(json.load(f))
        except Exception as e:
            logger.error(f"Error loading history file {file}: {str(e)}")
    
    # Sort by date (newest first)
    return sorted(history, key=lambda x: x.get('generatedDate', ''), reverse=True)

def create_docx_document(title, content_sections, metadata=None):
    """
    Creates a DOCX document with the given title, content sections, and metadata.
    
    Args:
        title (str): Document title
        content_sections (list): List of dictionaries with 'heading' and 'content' keys
        metadata (dict, optional): Document metadata (case number, date, etc.)
    
    Returns:
        BytesIO: Document as bytes stream
    """
    doc = docx.Document()
    
    # Set document margins (1 inch on all sides)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Add title
    title_paragraph = doc.add_paragraph()
    title_run = title_paragraph.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(16)
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add metadata if provided
    if metadata:
        metadata_paragraph = doc.add_paragraph()
        metadata_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for key, value in metadata.items():
            if value:
                metadata_paragraph.add_run(f"{key}: {value}\n")
        
        # Add spacing after metadata
        doc.add_paragraph()
    
    # Add content sections
    for section in content_sections:
        if section.get('heading'):
            heading = doc.add_heading(section['heading'], level=2)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        if section.get('content'):
            content = section['content']
            
            # Handle content based on type
            if isinstance(content, list):
                # If content is a list, create bullet points
                for item in content:
                    p = doc.add_paragraph(style='List Bullet')
                    p.add_run(item)
            else:
                # Plain text content
                paragraphs = content.split('\n')
                for para in paragraphs:
                    if para.strip():
                        p = doc.add_paragraph()
                        p.add_run(para)
    
    # Save document to memory stream
    docx_bytes = io.BytesIO()
    doc.save(docx_bytes)
    docx_bytes.seek(0)
    
    return docx_bytes

#############################
# Transcription Module Code #
#############################

def extract_audio(uploaded_file):
    """Extract audio from video file or process audio file directly"""
    input_path = None 
    try:
        if uploaded_file.name.lower().endswith(('.mp4', '.avi', '.mov')):
            if not check_ffmpeg():
                st.error("FFmpeg не установлен...")
                return None
        with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(uploaded_file.name)[1]) as tmpfile:
            tmpfile.write(uploaded_file.getbuffer())
            input_path = tmpfile.name
        if uploaded_file.name.lower().endswith(('.mp4', '.avi', '.mov')):
            audio_path = os.path.splitext(input_path)[0] + '.mp3'
            subprocess.run(["ffmpeg", "-i", input_path, "-q:a", "0", "-map", "a", audio_path], check=True, capture_output=True)
            if input_path and os.path.exists(input_path):
                 os.remove(input_path) # Удаляем видео
            return audio_path # Возвращаем mp3
        else:
            return input_path # Возвращаем аудио
    except Exception as e:
        st.error(f"Ошибка при извлечении аудио: {str(e)}")
        if input_path and os.path.exists(input_path): 
            os.remove(input_path) # Удаляем при ошибке
        return None # Возвращаем None

def transcribe_audio(client, audio_file, language='ru'):
    """Transcribe audio file using OpenAI Whisper"""
    try:
        with open(audio_file, "rb") as file:
            transcript = client.audio.transcriptions.create(
                model="whisper-1",
                file=file,
                language=language
            )
        return transcript.text
    except Exception as e:
        st.error(f"Ошибка при транскрибации: {str(e)}")
        return None
    finally:
        if os.path.exists(audio_file):
            os.remove(audio_file)

def analyze_transcription(client, text, analysis_type, language='ru'):
    """Analyze transcription text based on specified analysis type"""
    # System prompts based on analysis type
    prompts = {
        "summary": f"Вы опытный следователь. Суммируйте следующий текст показаний на языке {language}, выделив ключевую информацию:",
        "sequence": "Вы следователь, оценивающий последовательность показаний. Проанализируйте текст и выявите нарушения логической последовательности или пропущенные детали:",
        "facts": "Вы следователь, выделяющий существенные факты. Извлеките из текста ключевые факты, имеющие значение для следствия, в виде списка:",
    }
    
    try:
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": prompts.get(analysis_type, prompts["summary"])},
                {"role": "user", "content": text}
            ],
            temperature=0.5,
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        st.error(f"Ошибка при анализе: {str(e)}")
        return None

def compare_testimonies(client, text1, text2):
    """Compare two testimonies to find contradictions and extract relevant quotes.
    Returns a list of dictionaries, each representing a contradiction.
    """
    prompt = (
        "Вы следователь, сопоставляющий показания для выявления противоречий. "
        "Сравните следующие два показания и определите существенные противоречия между ними. "
        "Для каждого найденного противоречия предоставьте:\n"
        "1. Краткое описание сути противоречия.\n"
        "2. Точную цитату из показаний Лица №1, иллюстрирующую это противоречие.\n"
        "3. Точную цитату из показаний Лица №2, иллюстрирующую это противоречие.\n"
        "4. Оценку значимости противоречия (например, Низкая, Средняя, Высокая).\n\n"
        "Ответ должен быть представлен СТРОГО в формате JSON списка объектов, где каждый объект имеет ключи: "
        "'description' (строка), 'quote1' (строка), 'quote2' (строка), 'significance' (строка).\n"
        "Пример объекта JSON:\n"
        "{\n"
        "  \"description\": \"Время ухода из дома\",\n"
        "  \"quote1\": \"Я ушел из дома около 9 утра.\",\n"
        "  \"quote2\": \"Он вышел не раньше 11 часов.\",\n"
        "  \"significance\": \"Средняя\"\n"
        "}\n\n"
        "Если противоречий не найдено, верните пустой JSON список [].\n\n"
        "Показания лица №1:\n" + text1 + "\n\n"
        "Показания лица №2:\n" + text2
    )
    try:
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "Вы опытный следователь, точно извлекающий противоречия и цитаты из показаний в формате JSON."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.3,
            response_format={ "type": "json_object"} 
        )
        raw_response = response.choices[0].message.content.strip()
        try:
            contradictions_list = json.loads(raw_response)
            if isinstance(contradictions_list, list):
                return contradictions_list
            else:
                 if isinstance(contradictions_list, dict):
                     for key, value in contradictions_list.items():
                         if isinstance(value, list):
                             st.warning(f"Модель вернула JSON-объект вместо списка. Использую список из ключа '{key}'.")
                             return value
                 st.error("Модель вернула корректный JSON, но не в формате списка.")
                 print(f"Неожиданный JSON от OpenAI: {raw_response}")
                 return [] 
        except json.JSONDecodeError as json_e:
            st.error(f"Ошибка декодирования JSON ответа от OpenAI: {json_e}")
            st.warning("Модель не вернула валидный JSON. Противоречия не будут извлечены.")
            print(f"Невалидный JSON от OpenAI: {raw_response}")
            return [] 
    except Exception as e:
        st.error(f"Ошибка при сравнении показаний: {str(e)}")
        return None 

def generate_questions(client, contradictions):
    """Generate questions based on contradictions"""
    prompt = (
        "На основе следующих противоречий, выявленных в показаниях, сформулируйте список конкретных вопросов для уточнения "
        "и устранения противоречий:\n\n" + contradictions
    )
    
    try:
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "Вы опытный следователь, формулирующий точные вопросы для устранения противоречий в показаниях."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.7,
        )
        result = response.choices[0].message.content.strip()
        
        # Ensure result is in list format
        if not result.startswith("1.") and not result.startswith("-"):
            lines = result.split("\n")
            formatted_result = ""
            for i, line in enumerate(lines, 1):
                if line.strip():
                    formatted_result += f"{i}. {line.strip()}\n"
            return formatted_result
        
        return result
    except Exception as e:
        st.error(f"Ошибка при формировании вопросов: {str(e)}")
        return None

##########################
# Planning Module Code #
##########################

def extract_case_facts(client, case_description):
    """Extract key facts from case description"""
    try:
        prompt = (
            "Проанализируйте следующую фабулу дела и выделите ключевые факты, обстоятельства, "
            "доказательства и участников. Представьте информацию в виде структурированного списка:\n\n" + 
            case_description
        )
        
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "Вы опытный следователь, специализирующийся на анализе материалов дел."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.3,
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        st.error(f"Ошибка при анализе фабулы дела: {str(e)}")
        return None

def determine_crime_classification(client, facts):
    """Determine crime classification and relevant legal articles"""
    try:
        prompt = (
            "Проанализируйте следующие факты дела и определите:\n"
            "1. Квалификацию преступления (тип и категорию)\n"
            "2. Соответствующие статьи УК РК\n"
            "Представьте ответ в формате: Квалификация: [квалификация], Статьи: [список статей]\n\n" +
            facts
        )
        
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "Вы опытный юрист, специализирующийся на квалификации преступлений по законодательству Республики Казахстан."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.3,
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        st.error(f"Ошибка при определении квалификации: {str(e)}")
        return None

def create_investigation_plan(client, facts, classification, methodology_text=None):
    """Create investigation plan based on facts, classification and methodology"""
    try:
        methodology_part = ""
        if methodology_text:
            methodology_part = f"\nПри составлении плана используйте следующую методику:\n{methodology_text}\n"
            
        prompt = (
            f"Составьте план расследования на основе следующих фактов и квалификации дела:{methodology_part}\n\n"
            f"ФАКТЫ:\n{facts}\n\n"
            f"КВАЛИФИКАЦИЯ:\n{classification}\n\n"
            "План должен включать:\n"
            "1. Версии расследования (не менее 2-3)\n"
            "2. Первоочередные следственные действия\n"
            "3. Последующие следственные действия\n"
            "4. Необходимые экспертизы\n"
            "5. Оперативно-розыскные мероприятия\n\n"
            "Для каждого действия укажите цель, ожидаемый результат и приоритет."
        )
        
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "Вы опытный следователь с многолетним опытом планирования расследований преступлений."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.5,
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        st.error(f"Ошибка при составлении плана: {str(e)}")
        return None

def process_methodology(client, uploaded_file):
    """Process methodology file and extract key points"""
    try:
        # Save file temporarily
        with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmpfile:
            tmpfile.write(uploaded_file.getbuffer())
            file_path = tmpfile.name
            
        # Extract text from PDF - this is a simplified version
        # In a real application, you would use a PDF parsing library
        prompt = f"Из этой методики расследования выделите ключевые рекомендации, алгоритмы действий и важные моменты, которые следует учитывать при планировании расследования."
        
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "Вы эксперт по методикам расследования преступлений."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.3,
        )
        
        return response.choices[0].message.content.strip()
    except Exception as e:
        st.error(f"Ошибка при обработке методики: {str(e)}")
        return None
    finally:
        if os.path.exists(file_path):
            os.remove(file_path)

##########################
# Indictment Module Code #
##########################

def generate_indictment(client, case_number, crime_description, suspect_info, evidence_list, additional_info=None):
    """Generate indictment based on case information"""
    try:
        # Create evidence text, handling both regular and file evidence
        evidence_items = []
        for evidence in evidence_list:
            evidence_text = f"- {evidence.get('type')}: {evidence.get('description')}"
            if evidence.get('fileName'):
                evidence_text += f" (Файл: {evidence.get('fileName')})"
            evidence_items.append(evidence_text)
        
        evidence_text = "\n".join(evidence_items)
        
        additional_part = ""
        if additional_info:
            additional_part = f"\nДополнительная информация:\n{additional_info}"
            
        prompt = (
            f"Составьте обвинительный акт по делу №{case_number} на основе следующей информации:\n\n"
            f"ОПИСАНИЕ ПРЕСТУПЛЕНИЯ:\n{crime_description}\n\n"
            f"ДАННЫЕ О ПОДОЗРЕВАЕМОМ:\n{suspect_info}\n\n"
            f"ДОКАЗАТЕЛЬСТВА:\n{evidence_text}{additional_part}\n\n"
            "Обвинительный акт должен включать:\n"
            "1. Вводную часть (данные о подозреваемом, квалификация деяния)\n"
            "2. Описательную часть (описание события преступления, его обстоятельств)\n"
            "3. Доказательственную часть (перечень и анализ доказательств)\n"
            "4. Юридическую квалификацию (правовая оценка деяния)\n"
            "5. Заключительную часть (процессуальные решения)"
        )
        
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "Вы опытный прокурор, специализирующийся на составлении обвинительных актов."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.5,
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        st.error(f"Ошибка при составлении обвинительного акта: {str(e)}")
        return None

def analyze_evidence(client, evidence_list, crime_description):
    """Analyze evidence in relation to crime description"""
    try:
        # Create evidence text, handling both regular and file evidence
        evidence_items = []
        for evidence in evidence_list:
            evidence_text = f"- {evidence.get('type')}: {evidence.get('description')}"
            if evidence.get('fileName'):
                evidence_text += f" (Файл: {evidence.get('fileName')})"
                # Include file content snippet if available
                if evidence.get('fileContent'):
                    snippet = evidence.get('fileContent')[:200] + "..." if len(evidence.get('fileContent', '')) > 200 else evidence.get('fileContent', '')
                    if snippet:
                        evidence_text += f"\n  Фрагмент содержимого: {snippet}"
            evidence_items.append(evidence_text)
        
        evidence_text = "\n".join(evidence_items)
        
        prompt = (
            f"Проанализируйте следующие доказательства в контексте данного преступления и оцените их значимость, "
            f"достаточность и вероятные противоречия:\n\n"
            f"ПРЕСТУПЛЕНИЕ:\n{crime_description}\n\n"
            f"ДОКАЗАТЕЛЬСТВА:\n{evidence_text}\n\n"
            "Анализ должен включать:\n"
            "1. Оценку каждого доказательства (включая файлы)\n"
            "2. Выявление пробелов в доказательственной базе\n"
            "3. Рекомендации по дополнительным доказательствам"
        )
        
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "Вы опытный юрист, специализирующийся на анализе доказательств в уголовных делах."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.5,
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        st.error(f"Ошибка при анализе доказательств: {str(e)}")
        return None

#####################
# Main Application #
#####################

def main():
    # Initialize
    load_css()
    create_directories()
    client = init_openai()
    
    # Check dependencies
    ffmpeg_available = check_ffmpeg()
    
    # Sidebar navigation
    with st.sidebar:
        st.title("🔍 Система помощника следователя")
        
        module = st.radio(
            "Выберите модуль:",
            ["Транскрибация показаний", "Планирование расследования", "Генератор обвинительных актов"]
        )
        
        if not ffmpeg_available and module == "Транскрибация показаний":
            st.error("FFmpeg не установлен. Для работы с видео и аудио необходимо установить FFmpeg.")
        
        if not client:
            st.error("OpenAI API недоступен. Пожалуйста, проверьте настройки API ключа.")
    
    # Main content
    if module == "Транскрибация показаний":
        show_transcription_module(client)
    elif module == "Планирование расследования":
        show_planning_module(client)
    elif module == "Генератор обвинительных актов":
        show_indictment_module(client)
    
    # Footer
    st.markdown("""
        <div class="footer">
            © 2025 Система помощника следователя | Версия 1.0.0
        </div>
    """, unsafe_allow_html=True)

##########################################
# Module Interface Functions
##########################################

def show_transcription_module(client):
    st.title("🎙️ Транскрибация следственных действий")
    
    # Create tabs for main functionality and history
    tab1, tab2 = st.tabs(["Новая транскрибация", "История транскрибаций"])
    
    with tab1:
        st.header("Загрузка и обработка аудио/видео")
        
        # Form for transcription
        with st.form("transcription_form"):
            case_number = st.text_input("Номер материала", value=generate_case_number(), disabled=True)
            
            language = st.selectbox(
                "Язык материала",
                options=[("ru", "🇷🇺 Русский"), ("kk", "🇰🇿 Қазақ тілі"), ("en", "🇬🇧 English")],
                format_func=lambda x: x[1]
            )
            
            st.subheader("Загрузка файлов")
            uploaded_file1 = st.file_uploader(
                "Показания лица №1", 
                type=["mp4", "avi", "mov", "mp3", "wav"],
                key="file1"
            )
            
            uploaded_file2 = st.file_uploader(
                "Показания лица №2 (необязательно)", 
                type=["mp4", "avi", "mov", "mp3", "wav"],
                key="file2"
            )
            
            st.subheader("Дополнительные параметры")
            description = st.text_area("Описание материала (необязательно)")
            
            col1, col2 = st.columns(2)
            with col1:
                analyze_sequence = st.checkbox("Анализировать логическую последовательность", value=True)
                extract_facts = st.checkbox("Извлекать ключевые факты", value=True)
            
            with col2:
                find_contradictions = st.checkbox("Искать противоречия", value=True)
                generate_questions_check = st.checkbox("Генерировать уточняющие вопросы", value=True)
            
            submit_button = st.form_submit_button("Начать обработку")
        
        if submit_button:
            if not uploaded_file1:
                st.error("Необходимо загрузить файл с показаниями лица №1")
                return
            
            with st.spinner("Обработка файлов..."):
                try:
                    transcription_results = {"id": case_number, "language": language[0], "generatedDate": datetime.datetime.now().isoformat()}
                    transcription_results["description"] = description
                    transcription_results["statements"] = []
                    
                    # Process first file
                    st.info("Обработка файла лица №1...")
                    audio_path1 = extract_audio(uploaded_file1)
                    transcription1 = transcribe_audio(client, audio_path1, language[0])
                    
                    statement1 = {
                        "witnessName": "Лицо №1",
                        "fileUrl": uploaded_file1.name,
                        "transcription": transcription1,
                        "summary": "",
                        "logicalAnalysis": "",
                        "keyFacts": []
                    }
                    
                    # Analyze first transcription
                    if transcription1:
                        if analyze_sequence:
                            statement1["logicalAnalysis"] = analyze_transcription(client, transcription1, "sequence", language[0])
                        
                        if extract_facts:
                            facts_text = analyze_transcription(client, transcription1, "facts", language[0])
                            statement1["keyFacts"] = [fact.strip() for fact in facts_text.split('\n') if fact.strip()]
                        
                        statement1["summary"] = analyze_transcription(client, transcription1, "summary", language[0])
                    
                    transcription_results["statements"].append(statement1)
                    
                    # Process second file if uploaded
                    transcription2 = None
                    if uploaded_file2:
                        st.info("Обработка файла лица №2...")
                        audio_path2 = extract_audio(uploaded_file2)
                        transcription2 = transcribe_audio(client, audio_path2, language[0])
                        
                        statement2 = {
                            "witnessName": "Лицо №2",
                            "fileUrl": uploaded_file2.name,
                            "transcription": transcription2,
                            "summary": "",
                            "logicalAnalysis": "",
                            "keyFacts": []
                        }
                        
                        # Analyze second transcription
                        if transcription2:
                            if analyze_sequence:
                                statement2["logicalAnalysis"] = analyze_transcription(client, transcription2, "sequence", language[0])
                            
                            if extract_facts:
                                facts_text = analyze_transcription(client, transcription2, "facts", language[0])
                                statement2["keyFacts"] = [fact.strip() for fact in facts_text.split('\n') if fact.strip()]
                            
                            statement2["summary"] = analyze_transcription(client, transcription2, "summary", language[0])
                        
                        transcription_results["statements"].append(statement2)
                    
                    # Compare testimonies if requested
                    transcription_results['contradictions'] = [] 
                    if find_contradictions and transcription1 and transcription2: 
                        st.info("Сравнение показаний и поиск противоречий...")
                        contradictions_result = compare_testimonies(client, transcription1, transcription2)
                        if contradictions_result is not None:
                            transcription_results['contradictions'] = contradictions_result # Сохраняем список
                            if not contradictions_result: st.info("Существенных противоречий не выявлено.")
                        else: st.error("Не удалось сравнить показания из-за ошибки API.")

                    # Generate questions if requested and contradictions found
                    transcription_results['suggestedQuestions'] = [] 
                    if generate_questions_check and transcription_results['contradictions']:
                        st.info("Генерация уточняющих вопросов...")
                        contradictions_text = "\n".join([f"- {c.get('description', '')}" for c in transcription_results['contradictions']])
                        questions = generate_questions(client, contradictions_text)
                        if questions:
                             question_items = [q.strip() for q in questions.split('\n') if q.strip()]
                             transcription_results["suggestedQuestions"] = question_items
                        else: st.error("Не удалось сгенерировать вопросы.")
                    
                    # Save results
                    save_history("transcriptions", transcription_results)
                    
                    # Display success message
                    st.success("Обработка завершена успешно!")
                    
                    # Create DOCX document
                    content_sections = []
                    
                    # Add statements
                    for i, statement in enumerate(transcription_results.get('statements', []), 1):
                        content_sections.append({
                            'heading': f"Показания лица #{i}",
                            'content': statement.get('transcription', '')
                        })
                        
                        if statement.get('summary'):
                            content_sections.append({
                                'heading': f"Краткое резюме показаний лица #{i}",
                                'content': statement.get('summary', '')
                            })
                            
                        if statement.get('keyFacts'):
                            content_sections.append({
                                'heading': f"Ключевые факты показаний лица #{i}",
                                'content': statement.get('keyFacts', [])
                            })
                    
                    # Add contradictions if any
                    if transcription_results.get('contradictions'):
                        contradictions_content = []
                        for contradiction in transcription_results.get('contradictions', []):
                            contradictions_content.append(contradiction.get('description', ''))
                        
                        content_sections.append({
                            'heading': "Выявленные противоречия",
                            'content': contradictions_content
                        })
                    
                    # Add questions if any
                    if transcription_results.get('suggestedQuestions'):
                        content_sections.append({
                            'heading': "Уточняющие вопросы",
                            'content': transcription_results.get('suggestedQuestions', [])
                        })
                    
                    # Create metadata
                    metadata = {
                        "Номер материала": transcription_results.get('id', ''),
                        "Дата обработки": transcription_results.get('generatedDate', '')[:10],
                        "Язык": transcription_results.get('language', '')
                    }
                    
                    docx_bytes = create_docx_document(
                        f"Протокол транскрибации {transcription_results.get('id', '')}",
                        content_sections,
                        metadata
                    )
                    
                    col1, col2 = st.columns(2)
                    
                    with col1:
                        st.download_button(
                            "📄 Скачать протокол (DOCX)",
                            data=docx_bytes,
                            file_name=f"transcription_{transcription_results['id']}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            key=f"download_docx_main_{transcription_results['id']}"
                        )
                    
                    with col2:
                        st.download_button(
                            "📊 Скачать данные (JSON)",
                            data=json.dumps(transcription_results, ensure_ascii=False, indent=2),
                            file_name=f"transcription_{transcription_results['id']}.json",
                            mime="application/json",
                            key=f"download_json_main_{transcription_results['id']}"
                        )
                    
                except Exception as e:
                    st.error(f"Ошибка при обработке: {str(e)}")
    
    with tab2:
        st.header("История транскрибаций")
        
        # Load history
        history = load_history("transcriptions")
        
        if not history:
            st.info("История транскрибаций пуста")
        else:
            for item in history:
                with st.expander(f"Материал {item['id']} от {item.get('generatedDate', 'N/A')[:10]}"):
                    st.write(f"**Язык:** {item.get('language', 'Не указан')}")
                    st.write(f"**Количество файлов:** {len(item.get('statements', []))}")
                    
                    # Заменяем вложенные expander на tabs
                    tab_titles = [f"Показания {s.get('witnessName', f'Лицо #{i+1}')}" for i, s in enumerate(item.get('statements', []))]
                    if tab_titles:
                        tabs = st.tabs(tab_titles)
                        for i, statement in enumerate(item.get('statements', [])):
                             with tabs[i]:
                                st.write(f"Файл: {statement.get('fileUrl', 'Не указан')}")
                                st.text_area("Транскрипция", statement.get('transcription', ''), height=200, key=f"hist_transcription_{item['id']}_{i}", label_visibility="collapsed")
                                
                                # Убираем вложенные expander для деталей
                                if statement.get('summary'):
                                    st.subheader("Краткое резюме")
                                    st.write(statement.get('summary', ''))
                                if statement.get('logicalAnalysis'):
                                    st.subheader("Анализ последовательности")
                                    st.write(statement.get('logicalAnalysis', ''))
                                if statement.get('keyFacts'):
                                    st.subheader("Ключевые факты")
                                    if isinstance(statement.get('keyFacts'), list):
                                        for fact in statement.get('keyFacts', []):
                                            st.markdown(f"- {fact}")
                                    else: st.write(statement.get('keyFacts', ''))
                    
                    # Show contradictions if any
                    if 'contradictions' in item and item['contradictions']:
                        st.subheader("Выявленные противоречия")
                        if isinstance(item['contradictions'], list):
                            for contradiction in item['contradictions']:
                                if isinstance(contradiction, dict):
                                    st.markdown(f"**- {contradiction.get('description', 'Нет описания')}** (Значимость: {contradiction.get('significance', 'Не указана')})")
                                    st.caption(f"  _Лицо 1:_ {contradiction.get('quote1', 'Цитата отсутствует')}")
                                    st.caption(f"  _Лицо 2:_ {contradiction.get('quote2', 'Цитата отсутствует')}")
                                elif isinstance(contradiction, str): st.markdown(f"- {contradiction}") 
                        elif isinstance(item['contradictions'], str): st.markdown(item['contradictions'])
                    
                    # Show questions if any
                    if item.get('suggestedQuestions'):
                        st.subheader("Уточняющие вопросы")
                        for i, question in enumerate(item.get('suggestedQuestions', []), 1):
                            st.markdown(f"{i}. {question}")
                    
                    col1, col2, col3 = st.columns(3)
                    with col1:
                        st.download_button(
                            "Скачать результаты (JSON)",
                            data=json.dumps(item, ensure_ascii=False, indent=2),
                            file_name=f"transcription_{item['id']}.json",
                            mime="application/json",
                            key=f"download_json_hist_{item['id']}"
                        )
                    
                    with col2:
                        # Create DOCX document for download
                        content_sections = []
                        
                        # Add statements
                        for i, statement in enumerate(item.get('statements', []), 1):
                            content_sections.append({
                                'heading': f"Показания лица #{i}",
                                'content': statement.get('transcription', '')
                            })
                            
                            if statement.get('summary'):
                                content_sections.append({
                                    'heading': f"Краткое резюме показаний лица #{i}",
                                    'content': statement.get('summary', '')
                                })
                                
                            if statement.get('keyFacts'):
                                content_sections.append({
                                    'heading': f"Ключевые факты показаний лица #{i}",
                                    'content': statement.get('keyFacts', [])
                                })
                        
                        # Add contradictions if any
                        if item.get('contradictions'):
                            contradictions_content = []
                            for contradiction in item.get('contradictions', []):
                                contradictions_content.append(contradiction.get('description', ''))
                            
                            content_sections.append({
                                'heading': "Выявленные противоречия",
                                'content': contradictions_content
                            })
                        
                        # Add questions if any
                        if item.get('suggestedQuestions'):
                            content_sections.append({
                                'heading': "Уточняющие вопросы",
                                'content': item.get('suggestedQuestions', [])
                            })
                        
                        # Create metadata
                        metadata = {
                            "Номер материала": item.get('id', ''),
                            "Дата обработки": item.get('generatedDate', '')[:10],
                            "Язык": item.get('language', '')
                        }
                        
                        docx_bytes = create_docx_document(
                            f"Протокол транскрибации {item.get('id', '')}",
                            content_sections,
                            metadata
                        )
                        
                        st.download_button(
                            "Скачать (DOCX)",
                            data=docx_bytes,
                            file_name=f"transcription_{item['id']}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            key=f"download_docx_hist_{item['id']}"
                        )
                    
                    with col3:
                        if st.button("Удалить", key=f"delete_transcription_{item['id']}"):
                            # Implement deletion logic here
                            st.warning("Функция удаления будет доступна в следующей версии")

def show_planning_module(client):
    st.title("📋 Планирование расследования")
    
    # Create tabs for main functionality and history
    tab1, tab2 = st.tabs(["Новый план", "История планов"])
    
    with tab1:
        st.header("Формирование плана расследования")
        
        # Form for investigation planning
        with st.form("planning_form"):
            case_number = st.text_input("Номер дела", help="Введите уникальный номер дела")
            
            case_description = st.text_area(
                "Описание фабулы дела", 
                height=200,
                help="Введите подробное описание обстоятельств дела"
            )
            
            crime_category = st.selectbox(
                "Категория преступления",
                options=[
                    ("", "-- Выберите категорию (необязательно) --"),
                    ("theft", "Кража"),
                    ("robbery", "Грабеж"),
                    ("fraud", "Мошенничество"),
                    ("murder", "Убийство"),
                    ("assault", "Нанесение телесных повреждений"),
                    ("drugTrafficking", "Незаконный оборот наркотиков"),
                    ("other", "Другое")
                ],
                format_func=lambda x: x[1]
            )
            
            suspect_info = st.text_area(
                "Информация о подозреваемом (если известна)",
                height=100
            )
            
            methodology_file = st.file_uploader(
                "Загрузить методику расследования (PDF)",
                type=["pdf"],
                help="Загрузите файл методики для улучшения рекомендаций"
            )
            
            use_methodology = st.checkbox("Использовать базу методик расследования", value=True)
            
            submit_button = st.form_submit_button("Сформировать план расследования")
        
        if submit_button:
            if not case_number or not case_description:
                st.error("Необходимо заполнить номер дела и описание фабулы")
                return
            
            with st.spinner("Анализ дела и формирование плана..."):
                try:
                    # Extract facts from case description
                    facts = extract_case_facts(client, case_description)
                    
                    # Determine crime classification
                    classification = determine_crime_classification(client, facts)
                    
                    # Process methodology if provided
                    methodology_text = None
                    if methodology_file and use_methodology:
                        st.info("Обработка методики расследования...")
                        methodology_text = process_methodology(client, methodology_file)
                    
                    # Create investigation plan
                    st.info("Формирование плана расследования...")
                    plan = create_investigation_plan(client, facts, classification, methodology_text)
                    
                    # Prepare results
                    planning_results = {
                        "id": case_number,
                        "caseNumber": case_number,
                        "crimeCategory": crime_category[0] if crime_category[0] else "Не указана",
                        "suspectInfo": suspect_info,
                        "caseDescription": case_description,
                        "extractedFacts": facts,
                        "crimeClassification": classification,
                        "generatedDate": datetime.datetime.now().isoformat(),
                        "plan": plan
                    }
                    
                    # Parse classification to extract legal articles
                    try:
                        legal_articles = []
                        if "Статьи:" in classification:
                            articles_part = classification.split("Статьи:")[1].strip()
                            legal_articles = [art.strip() for art in articles_part.split(",")]
                        planning_results["legalArticles"] = legal_articles
                    except:
                        planning_results["legalArticles"] = []
                    
                    # Add methodology information
                    if methodology_text:
                        planning_results["methodologyUsed"] = True
                        planning_results["methodologyReferences"] = ["Загруженная пользователем методика"]
                    else:
                        planning_results["methodologyUsed"] = use_methodology
                        planning_results["methodologyReferences"] = []
                    
                    # Save results
                    save_history("planning", planning_results)
                    
                    # Display results
                    st.success("План расследования успешно сформирован!")
                    
                    st.subheader("Извлеченные факты")
                    st.markdown(facts)
                    
                    st.subheader("Квалификация преступления")
                    st.markdown(classification)
                    
                    st.subheader("План расследования")
                    st.markdown(plan)
                    
                    # Create DOCX document
                    content_sections = []
                    
                    # Add case description
                    content_sections.append({
                        'heading': "Фабула дела",
                        'content': case_description
                    })
                    
                    # Add extracted facts
                    content_sections.append({
                        'heading': "Извлеченные факты",
                        'content': facts
                    })
                    
                    # Add crime classification
                    content_sections.append({
                        'heading': "Квалификация преступления",
                        'content': classification
                    })
                    
                    # Add investigation plan
                    content_sections.append({
                        'heading': "План расследования",
                        'content': plan
                    })
                    
                    # Add methodology information if any
                    if methodology_text:
                        content_sections.append({
                            'heading': "Использованная методика",
                            'content': "Загруженная пользователем методика"
                        })
                    
                    # Create metadata
                    metadata = {
                        "Номер дела": case_number,
                        "Дата формирования": datetime.datetime.now().strftime("%Y-%m-%d"),
                        "Категория": planning_results.get('crimeCategory', 'Не указана')
                    }
                    
                    docx_bytes = create_docx_document(
                        f"План расследования по делу {case_number}",
                        content_sections,
                        metadata
                    )
                    
                    col1, col2 = st.columns(2)
                    
                    with col1:
                        st.download_button(
                            "📄 Скачать план (DOCX)",
                            data=docx_bytes,
                            file_name=f"plan_{case_number}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            key=f"download_plan_docx_main_{case_number}"
                        )
                    
                    with col2:
                        st.download_button(
                            "📊 Скачать данные (JSON)",
                            data=json.dumps(planning_results, ensure_ascii=False, indent=2),
                            file_name=f"plan_{case_number}.json",
                            mime="application/json",
                            key=f"download_plan_json_main_{case_number}"
                        )
                    
                except Exception as e:
                    st.error(f"Ошибка при формировании плана: {str(e)}")
    
    with tab2:
        st.header("История планов расследования")
        
        # Load history
        history = load_history("planning")
        
        if not history:
            st.info("История планов пуста")
        else:
            for item in history:
                with st.expander(f"Дело {item.get('caseNumber', 'N/A')} от {item.get('generatedDate', 'N/A')[:10]}"):
                    # Заменяем вложенные expander на tabs
                    tab1, tab2, tab3, tab4 = st.tabs(["Описание дела", "Извлеченные факты", "Квалификация", "План расследования"])
                    with tab1: st.write(item.get('caseDescription', ''))
                    with tab2: st.write(item.get('extractedFacts', ''))
                    with tab3: st.write(item.get('crimeClassification', ''))
                    with tab4: st.write(item.get('plan', ''))
                    
                    col1, col2, col3 = st.columns(3)
                    with col1:
                        st.download_button(
                            "Скачать план (JSON)",
                            data=json.dumps(item, ensure_ascii=False, indent=2),
                            file_name=f"plan_{item.get('caseNumber', 'case')}.json",
                            mime="application/json",
                            key=f"download_plan_json_hist_{item.get('id', '')}"
                        )
                    
                    with col2:
                        # Create DOCX document for download
                        content_sections = []
                        
                        # Add case description
                        content_sections.append({
                            'heading': "Фабула дела",
                            'content': item.get('caseDescription', '')
                        })
                        
                        # Add extracted facts
                        content_sections.append({
                            'heading': "Извлеченные факты",
                            'content': item.get('extractedFacts', '')
                        })
                        
                        # Add crime classification
                        content_sections.append({
                            'heading': "Квалификация преступления",
                            'content': item.get('crimeClassification', '')
                        })
                        
                        # Add investigation plan
                        content_sections.append({
                            'heading': "План расследования",
                            'content': item.get('plan', '')
                        })
                        
                        # Add methodology references if any
                        if item.get('methodologyReferences'):
                            content_sections.append({
                                'heading': "Использованные методики",
                                'content': item.get('methodologyReferences', [])
                            })
                        
                        # Create metadata
                        metadata = {
                            "Номер дела": item.get('caseNumber', ''),
                            "Дата формирования": item.get('generatedDate', '')[:10],
                            "Категория": item.get('crimeCategory', '')
                        }
                        
                        docx_bytes = create_docx_document(
                            f"План расследования по делу {item.get('caseNumber', '')}",
                            content_sections,
                            metadata
                        )
                        
                        st.download_button(
                            "Скачать (DOCX)",
                            data=docx_bytes,
                            file_name=f"plan_{item.get('caseNumber', '')}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            key=f"download_plan_docx_hist_{item.get('id', '')}"
                        )
                    
                    with col3:
                        if st.button("Удалить", key=f"delete_planning_{item.get('id', '')}"):
                            # Implement deletion logic here
                            st.warning("Функция удаления будет доступна в следующей версии")

def show_indictment_module(client):
    st.title("🧑‍⚖️ Генератор обвинительных актов")
    
    # Create tabs for main functionality and history
    tab1, tab2 = st.tabs(["Новый акт", "История актов"])
    
    with tab1:
        st.header("Формирование обвинительного акта")
        
        # Form for indictment generation
        with st.form("indictment_form"):
            case_number = st.text_input("Номер дела", help="Введите уникальный номер дела")
            
            crime_description = st.text_area(
                "Описание преступления", 
                height=200,
                help="Введите подробное описание события преступления, включая время, место, способ совершения и последствия"
            )
            
            suspect_info = st.text_area(
                "Данные о подозреваемом",
                height=150,
                help="Укажите ФИО, дату рождения, место жительства, гражданство, образование, семейное положение, место работы, судимость"
            )
            
            st.subheader("Доказательства")
            st.info("Добавьте информацию о доказательствах по делу")
            
            # File uploader for evidence files
            uploaded_evidence_files = st.file_uploader(
                "Загрузите файлы доказательств", 
                type=["txt", "pdf", "docx", "doc"],
                accept_multiple_files=True,
                help="Поддерживаемые форматы: TXT, PDF, DOCX, DOC"
            )
            
            # Process uploaded files if any
            evidence_from_files = []
            if uploaded_evidence_files:
                st.subheader("Загруженные файлы доказательств")
                for file in uploaded_evidence_files:
                    # Create a unique file reference ID
                    file_ref = f"{datetime.datetime.now().strftime('%Y%m%d%H%M%S')}_{file.name}"
                    
                    # Store file content
                    file_content = ""
                    try:
                        if file.type == "text/plain":
                            # For txt files
                            file_content = file.getvalue().decode("utf-8")
                        elif file.type == "application/pdf":
                            # For PDF files - in a real app, you'd use a PDF parser here
                            file_content = f"[PDF файл: {file.name}]"
                        else:
                            # For other document types
                            file_content = f"[Документ: {file.name}]"
                        
                        # Save file to storage
                        create_directories()
                        file_path = f"storage/evidence/{file_ref}"
                        with open(file_path, "wb") as f:
                            f.write(file.getbuffer())
                        
                        # Create evidence item from file
                        col1, col2 = st.columns([1, 3])
                        with col1:
                            file_evidence_type = st.selectbox(
                                f"Тип доказательства ({file.name})",
                                options=[
                                    "Показания свидетеля", "Показания потерпевшего", "Показания подозреваемого",
                                    "Заключение эксперта", "Протокол осмотра", "Протокол обыска", "Иное"
                                ],
                                key=f"file_evidence_type_{file.name}"
                            )
                        
                        with col2:
                            # If it's a text file, show preview
                            if file.type == "text/plain":
                                with st.expander(f"Просмотр содержимого файла {file.name}"):
                                    st.text(file_content[:2000] + ("..." if len(file_content) > 2000 else ""))
                            
                            # Allow user to add a description
                            file_evidence_description = st.text_area(
                                f"Описание доказательства ({file.name})",
                                height=100,
                                key=f"file_evidence_description_{file.name}"
                            )
                        
                        # Add to evidence list
                        evidence_from_files.append({
                            "type": file_evidence_type,
                            "description": file_evidence_description if file_evidence_description else f"Файл: {file.name}",
                            "fileReference": file_ref,
                            "fileName": file.name,
                            "fileContent": file_content[:5000] if file.type == "text/plain" else ""
                        })
                    
                    except Exception as e:
                        st.error(f"Ошибка при обработке файла {file.name}: {str(e)}")
            
            st.subheader("Ручное добавление доказательств")
            
            # Initialize session state for evidence list
            if 'evidence_list' not in st.session_state:
                st.session_state.evidence_list = [{"type": "", "description": ""}]
            
            # Display evidence items for manual entry
            evidence_list = []
            for i, evidence in enumerate(st.session_state.evidence_list):
                col1, col2 = st.columns([1, 3])
                with col1:
                    evidence_type = st.selectbox(
                        "Тип доказательства",
                        options=[
                            "", "Показания свидетеля", "Показания потерпевшего", "Показания подозреваемого",
                            "Заключение эксперта", "Протокол осмотра", "Протокол обыска", "Иное"
                        ],
                        key=f"evidence_type_{i}",
                        index=0 if evidence["type"] == "" else ["", "Показания свидетеля", "Показания потерпевшего", 
                                                             "Показания подозреваемого", "Заключение эксперта", 
                                                             "Протокол осмотра", "Протокол обыска", "Иное"].index(evidence["type"])
                    )
                
                with col2:
                    evidence_description = st.text_area(
                        "Описание доказательства",
                        value=evidence["description"],
                        height=100,
                        key=f"evidence_description_{i}"
                    )
                
                evidence_list.append({"type": evidence_type, "description": evidence_description})
            
            # Combine both manual and file-based evidence
            combined_evidence_list = evidence_from_files + evidence_list
            
            # Add/remove evidence buttons
            col1, col2 = st.columns(2)
            with col1:
                if st.form_submit_button("➕ Добавить доказательство"):
                    st.session_state.evidence_list.append({"type": "", "description": ""})
                    st.rerun()
            
            with col2:
                if len(st.session_state.evidence_list) > 1 and st.form_submit_button("➖ Удалить последнее"):
                    st.session_state.evidence_list.pop()
                    st.rerun()
            
            additional_info = st.text_area(
                "Дополнительная информация",
                height=100,
                help="Укажите любую дополнительную информацию, которая может быть полезна при составлении акта"
            )
            
            submit_button = st.form_submit_button("Сформировать обвинительный акт")
        
        if submit_button:
            if not case_number or not crime_description or not suspect_info:
                st.error("Необходимо заполнить номер дела, описание преступления и данные о подозреваемом")
                return
            
            # Validate evidence (from combined list)
            valid_evidence = [e for e in combined_evidence_list if e.get("type") and e.get("description")]
            if not valid_evidence:
                st.error("Необходимо добавить хотя бы одно доказательство")
                return
            
            with st.spinner("Формирование обвинительного акта..."):
                try:
                    # Analyze file evidence content if available
                    file_evidence_content = ""
                    for evidence in valid_evidence:
                        if evidence.get("fileContent"):
                            file_evidence_content += f"\nИз файла '{evidence.get('fileName', '')}':\n{evidence.get('fileContent')}\n"
                    
                    # Add file content to additional info if available
                    enhanced_additional_info = additional_info
                    if file_evidence_content:
                        enhanced_additional_info = (additional_info or "") + "\n\nСодержимое файлов доказательств:\n" + file_evidence_content
                    
                    # Generate indictment
                    indictment_text = generate_indictment(
                        client, 
                        case_number, 
                        crime_description, 
                        suspect_info, 
                        valid_evidence, 
                        enhanced_additional_info
                    )
                    
                    # Analyze evidence
                    evidence_analysis = analyze_evidence(client, valid_evidence, crime_description)
                    
                    # Extract defendant name from suspect info
                    defendant = "Подозреваемый"
                    suspect_lines = suspect_info.split('\n')
                    if suspect_lines and ":" in suspect_lines[0]:
                        defendant = suspect_lines[0].split(':')[1].strip()
                    elif len(suspect_lines) > 0:
                        defendant = suspect_lines[0].strip()
                    
                    # Prepare results
                    indictment_results = {
                        "id": case_number,
                        "caseNumber": case_number,
                        "defendant": defendant,
                        "crimeDescription": crime_description,
                        "suspectInfo": suspect_info,
                        "evidenceList": valid_evidence,
                        "additionalInformation": additional_info,
                        "indictmentText": indictment_text,
                        "evidenceAnalysis": evidence_analysis,
                        "generatedDate": datetime.datetime.now().isoformat()
                    }
                    
                    # Parse indictment to extract parts
                    try:
                        parts = indictment_text.split("##")
                        indictment_parts = {}
                        current_part = ""
                        for part in parts:
                            if part.strip():
                                lines = part.strip().split('\n')
                                if lines and lines[0].strip():
                                    part_name = lines[0].strip().lower()
                                    if "вводная" in part_name:
                                        indictment_parts["introductionText"] = '\n'.join(lines[1:]).strip()
                                    elif "описательная" in part_name:
                                        indictment_parts["descriptionText"] = '\n'.join(lines[1:]).strip()
                                    elif "доказательств" in part_name:
                                        indictment_parts["evidenceAnalysisText"] = '\n'.join(lines[1:]).strip()
                                    elif "заключение" in part_name or "заключительная" in part_name:
                                        indictment_parts["conclusionText"] = '\n'.join(lines[1:]).strip()
                                    else:
                                        current_part = '\n'.join(lines).strip()
                                else:
                                    current_part = part.strip()
                        
                        indictment_results.update(indictment_parts)
                        
                        # If parts weren't parsed correctly, use the full text
                        if not indictment_parts:
                            indictment_results["introductionText"] = "Смотрите полный текст"
                            indictment_results["descriptionText"] = "Смотрите полный текст"
                            indictment_results["evidenceAnalysisText"] = "Смотрите полный текст"
                            indictment_results["conclusionText"] = "Смотрите полный текст"
                    except:
                        indictment_results["introductionText"] = "Ошибка при разборе текста"
                        indictment_results["descriptionText"] = "Ошибка при разборе текста"
                        indictment_results["evidenceAnalysisText"] = "Ошибка при разборе текста"
                        indictment_results["conclusionText"] = "Смотрите полный текст"
                    
                    # Save results
                    save_history("indictments", indictment_results)
                    
                    # Display results
                    st.success("Обвинительный акт успешно сформирован!")
                    
                    st.subheader("Обвинительный акт")
                    st.markdown(indictment_text)
                    
                    st.subheader("Анализ доказательств")
                    st.markdown(evidence_analysis)
                    
                    # Create DOCX document
                    content_sections = []
                    sections_found = False 
                    section_map = {
                        "ВВОДНАЯ ЧАСТЬ": "introductionText",
                        "ОПИСАТЕЛЬНАЯ ЧАСТЬ": "descriptionText",
                        "ДОКАЗАТЕЛЬСТВА": "evidenceAnalysisText",
                        "ЗАКЛЮЧЕНИЕ": "conclusionText"
                    }
                    for heading, key in section_map.items():
                        content = indictment_results.get(key, '')
                        if content and content != "Смотрите полный текст":
                            content_sections.append({'heading': heading, 'content': content})
                            sections_found = True
                    if not sections_found:
                        st.warning("Не удалось разделить текст акта на секции. DOCX будет содержать полный текст.")
                        content_sections = [{'heading': "", 'content': indictment_text}]
                    
                    # Create metadata
                    metadata = {
                        "Номер дела": case_number,
                        "Дата формирования": datetime.datetime.now().strftime("%Y-%m-%d"),
                        "Обвиняемый": defendant
                    }
                    
                    docx_bytes = create_docx_document(
                        "ОБВИНИТЕЛЬНЫЙ АКТ",
                        content_sections,
                        metadata
                    )
                    
                    col1, col2, col3 = st.columns(3)
                    
                    with col1:
                        st.download_button(
                            "📄 Скачать акт (DOCX)",
                            data=docx_bytes,
                            file_name=f"indictment_{case_number}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            key=f"download_indict_docx_main_{case_number}"
                        )
                    
                    with col2:
                        st.download_button(
                            "📝 Скачать акт (текст)",
                            data=indictment_text,
                            file_name=f"indictment_{case_number}.txt",
                            mime="text/plain",
                            key=f"download_indict_txt_main_{case_number}"
                        )
                        
                    with col3:
                        st.download_button(
                            "📊 Скачать данные (JSON)",
                            data=json.dumps(indictment_results, ensure_ascii=False, indent=2),
                            file_name=f"indictment_{case_number}.json",
                            mime="application/json",
                            key=f"download_indict_json_main_{case_number}"
                        )
                    
                except Exception as e:
                    st.error(f"Ошибка при формировании обвинительного акта: {str(e)}")
    
    with tab2:
        st.header("История обвинительных актов")
        
        # Load history
        history = load_history("indictments")
        
        if not history:
            st.info("История обвинительных актов пуста")
        else:
            for item in history:
                with st.expander(f"Дело {item.get('caseNumber', 'N/A')} - {item.get('defendant', 'Подсудимый')}"):
                    st.write(f"**Дата:** {item.get('generatedDate', 'N/A')[:10]}")
                    
                    # Заменяем вложенные expander на tabs
                    tab1, tab2, tab3, tab4 = st.tabs(["Описание", "Подозреваемый", "Доказательства", "Текст акта"])
                    with tab1: st.write(item.get('crimeDescription', ''))
                    with tab2: st.write(item.get('suspectInfo', ''))
                    with tab3:
                        if isinstance(item.get('evidenceList'), list):
                           for evidence in item.get('evidenceList', []):
                               evidence_text = f"**{evidence.get('type', '')}:** {evidence.get('description', '')}"
                               if evidence.get('fileName'): evidence_text += f" (Файл: {evidence.get('fileName')})"
                               st.markdown(evidence_text)
                        else: st.write(item.get('evidenceList', ''))
                    with tab4: st.write(item.get('indictmentText', ''))
                    
                    col1, col2, col3 = st.columns(3)
                    with col1:
                        st.download_button(
                            "Скачать акт (текст)",
                            data=item.get('indictmentText', ''),
                            file_name=f"indictment_{item.get('caseNumber', 'case')}.txt",
                            mime="text/plain",
                            key=f"download_indict_txt_hist_{item.get('id', '')}"
                        )
                    
                    with col2:
                        # Create DOCX document for download
                        content_sections = []
                        sections_found = False
                        section_map = {
                            "ВВОДНАЯ ЧАСТЬ": "introductionText",
                            "ОПИСАТЕЛЬНАЯ ЧАСТЬ": "descriptionText",
                            "ДОКАЗАТЕЛЬСТВА": "evidenceAnalysisText",
                            "ЗАКЛЮЧЕНИЕ": "conclusionText"
                        }
                        for heading, key in section_map.items():
                            content = item.get(key, '')
                            if content and content != "Смотрите полный текст":
                                content_sections.append({'heading': heading, 'content': content})
                                sections_found = True
                        if not sections_found:
                            content_sections = [{'heading': "", 'content': item.get('indictmentText', '')}]
                        
                        # Create metadata
                        metadata = {
                            "Номер дела": item.get('caseNumber', ''),
                            "Дата формирования": item.get('generatedDate', '')[:10],
                            "Обвиняемый": item.get('defendant', '')
                        }
                        
                        docx_bytes = create_docx_document(
                            "ОБВИНИТЕЛЬНЫЙ АКТ",
                            content_sections,
                            metadata
                        )
                        
                        st.download_button(
                            "Скачать (DOCX)",
                            data=docx_bytes,
                            file_name=f"indictment_{item.get('caseNumber', '')}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            key=f"download_indict_docx_hist_{item.get('id', '')}"
                        )
                    
                    with col3:
                        if st.button("Удалить", key=f"delete_indictment_{item.get('id', '')}"):
                            # Implement deletion logic here
                            st.warning("Функция удаления будет доступна в следующей версии")

if __name__ == "__main__":
    main()
